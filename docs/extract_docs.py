import docx
import os
import sys

docs_dir = os.path.dirname(os.path.abspath(__file__))
files = [
    "CodeAtlas_PRD_Document_1.docx",
    "CodeAtlas_TRD_Document_2.docx",
    "CodeAtlas_Application_Flow_Document_3.docx",
    "CodeAtlas_UI_UX_Design_Specification_Document_4.docx",
    "CodeAtlas_Backend_Schema_Document_5.docx",
    "CodeAtlas_Implementation_Plan_Document_6.docx",
    "CodeAtlas_Buildathon_2026_Idea_Proposal.docx",
]

for fname in files:
    fpath = os.path.join(docs_dir, fname)
    if not os.path.exists(fpath):
        print(f"--- FILE NOT FOUND: {fname} ---")
        continue
    doc = docx.Document(fpath)
    out_path = os.path.join(docs_dir, fname.replace(".docx", ".txt"))
    with open(out_path, "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")
        # Also extract tables
        for table in doc.tables:
            f.write("\n[TABLE]\n")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                f.write(" | ".join(cells) + "\n")
            f.write("[/TABLE]\n")
    print(f"Extracted: {fname} -> {os.path.basename(out_path)}")
